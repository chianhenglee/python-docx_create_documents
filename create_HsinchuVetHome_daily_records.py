
from docx import Document


def check_leap_year(input_year):
	if input_year % 4 == 0:
		if input_year % 100 == 0 :
			if input_year % 400 == 0 :
				LeapYear = True
			else :
				LeapYear = False
		else :
			LeapYear = True
	else :
		LeapYear = False

	return LeapYear

def get_mandarin_week(inp_num):
	'''return mandarin characters for week days (Sun,Mon,Tue...etc)'''
	if inp_num==0:
		return '日'
	if inp_num==1:
		return '一'
	if inp_num==2:
		return '二'
	if inp_num==3:
		return '三'
	if inp_num==4:
		return '四'
	if inp_num==5:
		return '五'
	if inp_num==6:
		return '六'


##### The main script starts here #####

inp_year = int(input('What year is it? Type and click enter: '))


print('What is the day on January 1st in '+str(inp_year)+'?')
inp_day = input('Sun=0, Mon=1, Tue=2 ...... Sat=6 : ')

print('Your input is: '+str(inp_year)+' and it is '+inp_day+' on January 1st.')

print('Start creating files......')



### Create a list of 365 or 366 int tuples (month,day)

if check_leap_year(inp_year):
	M2 = 29
	total_days = 366
	print('Leap Year!')
else:
	M2 = 28
	total_days = 365
	print('NOT a Leap Year!')

days_in_months = [31,M2,31,30,31,30,31,31,30,31,30,31]

date_list = []
for i in range(12):

	curr_month = [[inp_year,i+1,j] for j in range(1,days_in_months[i]+1)]
	date_list = date_list + curr_month
	#date_list.append(curr_month)



### Create a list of 365 or 366 numbers representing Sun Mon Tue Wed....Sat (0,1,2,3,4,5,6)
week_days_list = [];
s = int(inp_day)


for j in range(total_days):

	if j==0:
		curr_day = s
	else:
		if curr_day==6:
			curr_day = 0
		else:
			curr_day = curr_day+1


	week_days_list.append(curr_day)





### Now modify the template docx file and save as new file with the modified line and new filename ###
doc = Document('template_full.docx')
for i in range(total_days):

	print('Count:'+str(i+1))

	
	#print(doc.paragraphs[1].text)

	
	wd = get_mandarin_week(week_days_list[i])
	full_date_week = str(date_list[i][0])+'年'+str(date_list[i][1])+'月'+str(date_list[i][2])+'日 星期'+wd


	

	doc.paragraphs[1].text = '照護日期：'+full_date_week


	date_filename = str(date_list[i][0])+'.'+str(date_list[i][1])+'.'+str(date_list[i][2])
	doc.save(date_filename+'.docx')

	#print('New filed saved as .docx')

print('Done! A total of '+str(i+1)+' files created.')











