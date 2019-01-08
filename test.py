
from docx import Document



#import docx
doc = Document('template.docx')

print('Below we print the file line by line:')
print('')
for i in doc.paragraphs:
	print(i.text)

print('---- End of file ----')
print('')

print(doc.paragraphs[5].text)



doc.paragraphs[5].text = '這行是我新改的中文'

print(doc.paragraphs[5].text)



print('Below we print the new file content line by line:')
print('')
for i in doc.paragraphs:
	print(i.text)

print('---- End of file ----')
print('')

doc.save('template_modified.docx')
print('New filed saved as .docx')
